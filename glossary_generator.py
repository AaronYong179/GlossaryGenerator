import xlrd
import os
import tkinter as Tk
from tkinter import filedialog
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

class GlossaryGenerator:
    def __init__(self, filepath):
        self.glossary = {}
        self.sheet = xlrd.open_workbook(filepath).sheet_by_index(0)

    def _write(self, term, value, topic, page):
        page = int(page)
        term, value, topic, page = list(map(lambda x: str(x).strip(), (term, value, topic, page)))
        
        if term not in self.glossary:
            self.glossary[term] = {}
        if value not in self.glossary[term]:
            self.glossary[term][value] = []

        self.glossary[term][value].append((topic, page))

    def convert_glossary(self):
        for row in range(1, self.sheet.nrows):
            self._write(*self.sheet.row_values(row))


class GlossaryExporter:
    def __init__(self):
        self.document = Document()
        
    def format_docx(self, font_family):
        term_style = self.document.styles.add_style("Term", WD_STYLE_TYPE.PARAGRAPH)
        term_style.font.name = font_family
        term_style.paragraph_format.line_spacing = Inches(1/6)

        value_style = self.document.styles["List"]
        value_style.font.name = font_family
        value_style.paragraph_format.line_spacing = Inches(0.20)
        value_style.paragraph_format.left_indent = Inches(0.5)
        

    def export_glossary(self, filepath, glossary, font_family):
        self.format_docx(font_family)
        terms = sorted(list(glossary.keys()), key = str.casefold)
        for term in terms:
            values = sorted(list(glossary[term].keys()), key = str.casefold)
            self.document.add_paragraph(style = "Term").add_run(term).bold = True
            for value in values:
                topics_pgs = glossary[term][value]
                value_para = self.document.add_paragraph(value, style = "List")
                for topic_pg in topics_pgs:
                    value_para.add_run(f", ({topic_pg[0]}, {topic_pg[1]})")
                    
        self.document.save(filepath)


class MainApp:
    def __init__(self, parent):
        self.parent = parent
        self.run()

    def run(self):
        curr_dir = os.getcwd()
        input_dir = filedialog.askopenfilename(initialdir = curr_dir, title = "Please select a file to convert to a glossary.")
        output_dir = filedialog.asksaveasfilename(initialdir = curr_dir, title = "Please select a folder to save the glossary.", \
                                                  defaultextension = ".docx", filetypes = (("Word Document", "*.docx"),))  
        self.glossarygenerator = GlossaryGenerator(input_dir)
        self.glossarygenerator.convert_glossary()

        self.glossaryexporter = GlossaryExporter()
        self.glossaryexporter.export_glossary(output_dir, self.glossarygenerator.glossary, "Calibri")
        self.parent.destroy()




def main():
    root = Tk.Tk()
    MainApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()

# g = GlossaryGenerator("./2211_glossary.xlsx")
# g.convert_glossary()

# e = GlossaryExporter(".", "testing")
# e.export_glossary(g.glossary, "Calibri")


